"""
h11_manuscript.py — Generic manuscript generation handler (daemon-local).

Works for ANY project type: SSE, NMC, LiPS, LYC, DMB/electrolyte, Na-air, etc.
Project-specific content (title, abstract, intro, conclusion) is read from
the project YAML. Generic sections and equations are built from available data files.

Equations are inserted as proper Word OMML (Office Open XML Math, m:oMath),
which renders natively in Word's Insert → Equation environment using Cambria Math.
Never uses Courier-text approximations for math.

CLAUDE.md: every improvement here propagates to all future project runs.
"""
from __future__ import annotations

import json
import logging
import re
import sys
from pathlib import Path
from typing import TYPE_CHECKING

from .base import SimulationHandler
from hpca.core.combinations import is_combinatorial_parent
from hpca.core.paths import dft_opt, results_base, results_figures, results_manuscript, load_platform_config

if TYPE_CHECKING:
    from hpca.orchestrator.state_tracker import ProjectState

log = logging.getLogger("hpca.orch")
# Layout: see hpca/core/paths.py
# Cross-ref: hpca/core/paths.py results_manuscript(), results_figures(), results_base()
# HPC paths: hpca/config/platform.yaml hpc.cladue_site_packages

_CLADUE_SITE = load_platform_config().get("hpc", {}).get("cladue_site_packages", "")

# Optional global manuscript archive directory (set in platform.yaml hpc.global_manuscript_dir).
# Used by _next_version() to avoid repeating version numbers across projects.
# None means only the per-project ms_dir is searched.
_global_ms_str = load_platform_config().get("hpc", {}).get("global_manuscript_dir", "")
_GLOBAL_MS_DIR = Path(_global_ms_str) if _global_ms_str else None

# OMML namespace
_M = "http://schemas.openxmlformats.org/officeDocument/2006/math"
_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


# ═══════════════════════════════════════════════════════════════════════════════
# OMML builder — proper Word math equations
# ═══════════════════════════════════════════════════════════════════════════════

def _mk(tag: str):
    """Create an OMML element (uses lxml which python-docx already requires)."""
    from lxml import etree
    return etree.Element(f"{{{_M}}}{tag}")


def _text_run(text: str):
    """m:r/m:t element with plain text."""
    r = _mk("r")
    t = _mk("t")
    t.text = text
    r.append(t)
    return r


def _subscript(base_text: str, sub_text: str):
    """m:sSub: base_{sub}"""
    el = _mk("sSub")
    e  = _mk("e");   e.append(_text_run(base_text));  el.append(e)
    s  = _mk("sub"); s.append(_text_run(sub_text));   el.append(s)
    return el


def _superscript(base_text: str, sup_text: str):
    """m:sSup: base^{sup}"""
    el = _mk("sSup")
    e  = _mk("e");   e.append(_text_run(base_text));  el.append(e)
    s  = _mk("sup"); s.append(_text_run(sup_text));   el.append(s)
    return el


def _fraction(num_els: list, den_els: list):
    """m:f: num / den — each arg is a list of OMML elements."""
    f   = _mk("f")
    num = _mk("num"); [num.append(e) for e in num_els]; f.append(num)
    den = _mk("den"); [den.append(e) for e in den_els]; f.append(den)
    return f


def _build_omath(*parts) -> object:
    """
    Build an m:oMath element from a list of OMML child elements or plain strings.
    Plain strings are wrapped automatically in m:r/m:t.
    """
    omath = _mk("oMath")
    for p in parts:
        if isinstance(p, str):
            omath.append(_text_run(p))
        else:
            omath.append(p)
    return omath


def _insert_equation(doc, omath_el, center: bool = True) -> None:
    """
    Insert an m:oMath element as a proper Word equation paragraph.
    Word recognises m:oMath as a native equation (Insert → Equation style,
    rendered in Cambria Math) when placed directly inside w:p.
    """
    try:
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        p = doc.add_paragraph()
        if center:
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(6)
        # Append m:oMath directly inside w:p — Word treats this as a math zone
        p._p.append(omath_el)
    except Exception as exc:
        log.debug("[h11_manuscript] OMML insert failed: %s", exc)


# Pre-built equations for standard computational chemistry workflow
# Each is a function → call it to get a fresh oMath element (lxml elements are single-use)

def _eq_diffusivity():
    """Return OMML element for D = (1/6) dMSD/dt."""
    return _build_omath(
        "D = ",
        _fraction([_text_run("1")], [_text_run("6")]),
        " · ",
        _fraction(
            [_subscript("d⟨|Δr|", "2"), _text_run("⟩")],
            [_text_run("dt")],
        ),
    )


def _eq_arrhenius():
    """Return OMML element for the Arrhenius diffusivity equation D = D0 * exp(-Ea/kBT)."""
    return _build_omath(
        "D = ",
        _subscript("D", "0"),
        " · exp",
        _fraction(
            [_text_run("−"), _subscript("E", "a")],
            [_subscript("k", "B"), _text_run(" T")],
        ),
    )


def _eq_msd():
    """Return OMML element for MSD(τ) = <|r(t+τ) - r(t)|²>."""
    return _build_omath(
        "MSD(τ) = ⟨|r(t + τ) − r(t)|",
        _superscript("", "2"),
        "⟩",
    )


def _eq_nernst_einstein():
    """Return OMML element for the Nernst-Einstein ionic conductivity σ = nq²D/kBT."""
    return _build_omath(
        "σ = ",
        _fraction(
            [_text_run("n q"), _superscript("", "2"), _text_run(" D")],
            [_subscript("k", "B"), _text_run(" T")],
        ),
    )


def _eq_activation_energy():
    """Return OMML element for Ea = -kB × slope of ln D vs 1/T."""
    return _build_omath(
        _subscript("E", "a"),
        " = −",
        _subscript("k", "B"),
        " × slope of [ln D vs 1/T]",
    )


def _eq_koopmans_ip():
    """Return OMML element for IP ≈ -ε_HOMO (Koopmans' theorem)."""
    return _build_omath("IP ≈ −", _subscript("ε", "HOMO"))


def _eq_koopmans_ea():
    """Return OMML element for EA ≈ -ε_LUMO (Koopmans' theorem)."""
    return _build_omath("EA ≈ −", _subscript("ε", "LUMO"))


def _eq_vox():
    """Return OMML element for V_ox = IP - 1.37 V vs Li/Li+ (Trasatti reference)."""
    return _build_omath(
        _subscript("V", "ox"),
        " = IP − 1.37 V  (vs Li/Li",
        _superscript("", "+"),
        ",  Trasatti 1986)",
    )


def _eq_vred():
    """Return OMML element for V_red = EA - 1.37 V vs Li/Li+."""
    return _build_omath(
        _subscript("V", "red"),
        " = EA − 1.37 V  (vs Li/Li",
        _superscript("", "+"),
        ")",
    )


def _eq_rdf():
    """Return OMML element for the radial distribution function g(r) = <ρ(r)> / (4πr²ρ0)."""
    return _build_omath(
        "g(r) = ",
        _fraction(
            [_text_run("⟨ρ(r)⟩")],
            [_text_run("4πr"), _superscript("", "2"), _text_run("ρ"), _subscript("", "0")],
        ),
    )


def _eq_cn():
    """Return OMML element for the coordination number CN = 4πρ0 ∫ r² g(r) dr."""
    return _build_omath(
        "CN = 4π",
        _subscript("ρ", "0"),
        " ∫₀",
        _superscript("", "r_min"),
        " r",
        _superscript("", "2"),
        " g(r) dr",
    )


def _eq_neb_barrier():
    """Return OMML element for Ea = E_TS - E_IS (NEB migration barrier)."""
    return _build_omath(
        _subscript("E", "a"),
        " = ",
        _subscript("E", "TS"),
        " − ",
        _subscript("E", "IS"),
    )


# ═══════════════════════════════════════════════════════════════════════════════
# Handler
# ═══════════════════════════════════════════════════════════════════════════════

class ManuscriptHandler(SimulationHandler):
    """Daemon-local handler: generates versioned DOCX manuscript (any project type)."""

    name = "h11_manuscript"
    is_daemon = True

    _VARIANT_TITLES = {
        "cmd":      "Classical MD Study",
        "mlmd_dft": "DFT + Machine-Learning MD Study",
        "combined": "Combined CMD + DFT + MLMD Study",
    }

    def can_run(self, project_dir: Path, state: "ProjectState") -> bool:
        """Return True when at least one variant plot_manifest.json or legacy figure directory exists."""
        yaml = self.read_project_yaml(project_dir)
        # Combinatorial parent: wait until every sub-project manuscript is done
        if is_combinatorial_parent(yaml):
            sub_dirs = self._find_sub_project_dirs(project_dir, yaml)
            if not sub_dirs:
                return False
            return all(results_manuscript(d).exists() for d in sub_dirs)

        # Ready when any variant's plot_manifest.json exists
        for variant in ("cmd", "mlmd_dft", "combined"):
            if (results_figures(project_dir) / variant / "plot_manifest.json").exists():
                return True
        # Legacy fallback
        fig_dir = results_figures(project_dir)
        legacy_fig = project_dir / "Analysis" / "figures"
        return (fig_dir / "plot_manifest.json").exists() or \
               (fig_dir.is_dir() and len(list(fig_dir.glob("*.png"))) >= 1) or \
               (legacy_fig.is_dir() and len(list(legacy_fig.glob("*.png"))) >= 1)

    def is_complete(self, project_dir: Path, state: "ProjectState") -> bool:
        """Return True when all variant manuscript.docx files are newer than their plot_manifest.json."""
        any_variant = False
        for variant in ("cmd", "mlmd_dft", "combined"):
            manifest = results_figures(project_dir) / variant / "plot_manifest.json"
            if not manifest.exists():
                continue
            any_variant = True
            ms_path = project_dir / "results" / variant / "manuscript.docx"
            if not ms_path.exists():
                return False
            if ms_path.stat().st_mtime < manifest.stat().st_mtime:
                return False
        if any_variant:
            return True
        # Legacy fallback
        return results_manuscript(project_dir).exists()

    def submit(self, project_dir: Path, state: "ProjectState") -> str | None:
        """Generate DOCX manuscripts for each variant; delegates to combinatorial path for multi-composition projects."""
        yaml = self.read_project_yaml(project_dir)
        # Combinatorial parent → consolidated multi-composition manuscript
        if is_combinatorial_parent(yaml):
            return self._submit_combinatorial(project_dir, yaml, state)

        project_name  = yaml.get("name", project_dir.name)
        label_keywords = {
            "msd": ["msd", "mean_square"], "arrhenius": ["arrhenius", "activation"],
            "rdf": ["rdf", "radial"], "dfn": ["dfn", "doyle_fuller"],
        }
        any_generated = False

        for variant in ("cmd", "mlmd_dft", "combined"):
            manifest_path = results_figures(project_dir) / variant / "plot_manifest.json"
            if not manifest_path.exists():
                continue
            ms_dir  = project_dir / "results" / variant
            ms_path = ms_dir / "manuscript.docx"
            if ms_path.exists() and ms_path.stat().st_mtime >= manifest_path.stat().st_mtime:
                log.info("[h11_manuscript] variant=%s already up-to-date", variant)
                any_generated = True
                continue
            ms_dir.mkdir(parents=True, exist_ok=True)

            variant_fig_dir = results_figures(project_dir) / variant
            shared_fig_dir  = results_figures(project_dir)
            figure_paths: dict = {}
            for png in sorted(list(variant_fig_dir.glob("*.png")) +
                              list(shared_fig_dir.glob("*.png"))):
                stem = png.stem.lower()
                for label, kws in label_keywords.items():
                    if label not in figure_paths and any(k in stem for k in kws):
                        figure_paths[label] = png

            variant_title = f"{project_name} — {self._VARIANT_TITLES.get(variant, variant)}"
            log.info("[h11_manuscript] Generating %s", ms_path)
            generated = False
            try:
                from hpca.manuscript.generator import ManuscriptGenerator
                gen = ManuscriptGenerator(variant_title, ms_dir, yaml)
                pipeline_results = {
                    "project": variant_title, "params": yaml,
                    "analysis": {}, "continuum": {}, "benchmark": None,
                    "figures": figure_paths,
                }
                gen.generate_from_pipeline_run(pipeline_results)
                gen.save("manuscript.docx")
                generated = True
            except ImportError:
                pass
            except Exception as exc:
                log.warning("[h11_manuscript] [%s] generator failed: %s", variant, exc)

            if not generated:
                self._generate_generic_docx(project_dir, yaml, ms_path)

            log.info("[h11_manuscript] COMPLETE variant=%s: %s", variant, ms_path)
            any_generated = True

        if not any_generated:
            # Legacy fallback: single manuscript with flat Analysis/
            results_base(project_dir).mkdir(parents=True, exist_ok=True)
            out_path = results_manuscript(project_dir)
            figure_paths = {}
            for png in sorted(project_dir.glob("**/*.png")):
                stem = png.stem.lower()
                for label, kws in label_keywords.items():
                    if label not in figure_paths and any(k in stem for k in kws):
                        figure_paths[label] = png
            self._generate_generic_docx(project_dir, yaml, out_path)
            any_generated = True

        state.set_stage("h11_manuscript", "COMPLETE")
        return None

    # ── Combinatorial manuscript ─────────────────────────────────────────────

    def _find_sub_project_dirs(self, project_dir: Path, yaml_cfg: dict) -> list[Path]:
        """Return only the production children declared by the parent configuration."""
        from hpca.core.combinations import production_combinations
        return [project_dir / item["name"] for item in production_combinations(yaml_cfg)
                if (project_dir / item["name"] / "project.yaml").exists()]

    def _collect_sub_results(self, sub_dir: Path) -> dict:
        """Return a summary dict (name, Ea, echem, fig_dir) for one combinatorial sub-project."""
        r: dict = {"name": sub_dir.name, "dir": sub_dir}

        # Activation energy
        _arrh = self._find_arrhenius_csv(sub_dir)
        r["Ea_str"] = self._read_ea(_arrh) if _arrh else None

        # Echem window
        echem_json = sub_dir / "results" / "echem" / "echem_summary.json"
        r["echem"] = None
        if echem_json.exists():
            try:
                d = json.loads(echem_json.read_text())
                if "species" in d and d["species"]:
                    sp = d["species"][0]
                    r["echem"] = {"V_red": sp.get("V_red"), "V_ox": sp.get("V_ox")}
                elif "V_red" in d:
                    r["echem"] = {"V_red": d.get("V_red"), "V_ox": d.get("V_ox")}
            except Exception:
                pass

        # Figures directory
        r["fig_dir"] = results_figures(sub_dir)
        return r

    def _submit_combinatorial(self, project_dir: Path, yaml_cfg: dict,
                               state: "ProjectState") -> str | None:
        """Build a cross-composition DOCX manuscript for a multi-combination parent project."""
        sub_dirs = self._find_sub_project_dirs(project_dir, yaml_cfg)
        if not sub_dirs:
            log.warning("[h11_manuscript] No sub-project dirs found for combinatorial parent")
            return None

        results_base(project_dir).mkdir(parents=True, exist_ok=True)
        out_path = results_manuscript(project_dir)
        log.info("[h11_manuscript] Building combinatorial manuscript → %s", out_path)

        sub_results = [self._collect_sub_results(d) for d in sub_dirs]
        self._build_combinatorial_docx(project_dir, yaml_cfg, sub_results, out_path)

        state.set_stage("h11_manuscript", "COMPLETE", output=str(out_path))
        log.info("[h11_manuscript] COMPLETE (combinatorial): %s", out_path)
        return None

    def _build_combinatorial_docx(self, project_dir: Path, yaml_cfg: dict,
                                   sub_results: list[dict], out_path: Path) -> None:
        """Write the multi-composition DOCX with comparative table, per-composition sections, and shared references."""
        if _CLADUE_SITE not in sys.path:
            sys.path.insert(0, _CLADUE_SITE)

        try:
            from docx import Document
            from docx.shared import Pt, Inches, Cm, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
        except ImportError:
            log.error("[h11_manuscript] python-docx not available; writing text stub")
            lines = [f"Combinatorial Manuscript: {project_dir.name}\n"]
            for r in sub_results:
                lines.append(f"\n--- {r['name']} ---\nEa = {r.get('Ea_str', 'N/A')}\n")
            out_path.with_suffix(".txt").write_text("\n".join(lines))
            return

        doc = Document()
        sec = doc.sections[0]
        sec.top_margin = Cm(2.0);  sec.bottom_margin = Cm(2.0)
        sec.left_margin = Cm(2.5); sec.right_margin  = Cm(2.5)

        def _run(para, text, bold=False, italic=False, size=11, font="Times New Roman"):
            """Add a styled run to para and return it."""
            r = para.add_run(text)
            r.bold = bold; r.italic = italic
            r.font.name = font; r.font.size = Pt(size)
            return r

        def _h1(text):
            """Append a level-1 heading paragraph and return it."""
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after  = Pt(6)
            r = _run(p, text, bold=True, size=13); r.font.underline = True
            return p

        def _h2(text):
            """Append a level-2 heading paragraph and return it."""
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after  = Pt(4)
            _run(p, text, bold=True, size=11); return p

        def _body(text, size=11):
            """Append a body-text paragraph and return it."""
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(6)
            _run(p, text, size=size); return p

        def _add_table(headers, rows):
            """Append a styled Word table with a blue header row."""
            tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
            tbl.style = "Table Grid"
            hrow = tbl.rows[0]
            for i, h in enumerate(headers):
                cell = hrow.cells[i]; cell.text = ""
                r = cell.paragraphs[0].add_run(h)
                r.bold = True; r.font.size = Pt(9.5)
                r.font.color.rgb = RGBColor(255, 255, 255)
                tcPr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "2E75B6"); tcPr.append(shd)
            for ri, row in enumerate(rows):
                trow = tbl.rows[ri + 1]
                for ci, val in enumerate(row):
                    trow.cells[ci].text = ""
                    r = trow.cells[ci].paragraphs[0].add_run(str(val))
                    r.font.size = Pt(9.5)
            doc.add_paragraph()

        def _add_fig(fig_dir, glob_pat, caption):
            """Insert the first PNG matching glob_pat from fig_dir, or a placeholder paragraph."""
            fig = next(iter(sorted(fig_dir.glob(glob_pat))), None) if (fig_dir and fig_dir.is_dir()) else None
            if fig and fig.exists():
                try:
                    doc.add_picture(str(fig), width=Inches(5.5))
                    cap = doc.add_paragraph()
                    _run(cap, caption, italic=True, size=9.5)
                    cap.paragraph_format.space_after = Pt(10)
                    return
                except Exception:
                    pass
            _body(f"[Figure: {caption}]", size=9)

        # ── derive display names ──────────────────────────────────────────
        solvents = yaml_cfg.get("solvents", [])
        salts    = yaml_cfg.get("salts",    [])
        sol_str  = "/".join(solvents) if solvents else project_dir.name
        salt_str = "/".join(salts)    if salts    else ""

        # ── Title ─────────────────────────────────────────────────────────
        title_p = doc.add_paragraph(); title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title = yaml_cfg.get(
            "manuscript_title",
            f"Concentration-Dependent Transport and Electronic Properties of "
            f"{sol_str}{(' ' + salt_str) if salt_str else ''} Electrolytes: "
            "A Multi-Scale Computational Study",
        )
        _run(title_p, title, bold=True, size=14)

        auth_p = doc.add_paragraph(); auth_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(auth_p, yaml_cfg.get("authors", "Selva Chandrasekaran Selvaraj et al."), size=11)

        affil_p = doc.add_paragraph(); affil_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(affil_p,
             yaml_cfg.get("affiliations",
                          "National Renewable Energy Laboratory (NREL), Golden, CO 80401"),
             italic=True, size=10)
        doc.add_paragraph()

        # ── Abstract ──────────────────────────────────────────────────────
        _h1("Abstract")
        Ea_summary = "; ".join(
            f"{r['name']}: Ea = {r['Ea_str']}" for r in sub_results if r.get("Ea_str")
        ) or "multiple compositions investigated"
        _body(
            f"We present a systematic multi-scale computational study of {sol_str} electrolytes"
            f"{(' with ' + salt_str) if salt_str else ''} at varying salt concentrations "
            f"({', '.join(r['name'] for r in sub_results)}). "
            "Ab initio molecular dynamics (AIMD, VASP 6.4.2, PBE-GGA) training data was used "
            "to develop composition-specific DeepMD machine-learning interatomic potentials. "
            "Machine-learning molecular dynamics (MLMD, LAMMPS) at 300–700 K yields self-diffusion "
            "coefficients, ionic conductivities, and Arrhenius activation energies for each "
            f"composition. {Ea_summary}. "
            "Electronic structure (Bader charges, HOMO-LUMO) and solvation structure (RDF, CN) "
            "analyses reveal concentration-dependent structural and electronic changes."
        )

        # ── Introduction ──────────────────────────────────────────────────
        _h1("1. Introduction")
        _body(
            "Salt concentration is a key lever for tuning ionic transport, solvation structure, "
            "and electrochemical stability in liquid electrolytes. "
            f"This work systematically investigates {sol_str} electrolytes at "
            f"{len(sub_results)} concentrations using a unified multi-scale pipeline: "
            "DFT single-molecule optimisation, AIMD dataset generation, DeepMD potential training, "
            "and MLMD production simulations. The same protocol is applied to each composition, "
            "enabling a direct, consistent comparison of transport properties."
        )

        # ── Methods ───────────────────────────────────────────────────────
        _h1("2. Computational Methods")
        _body(
            "Identical simulation protocols were applied to all compositions. "
            "DFT geometry optimisations (VASP 6.4.2, PBE-D3, ENCUT=520 eV, EDIFF=1E-5) "
            "were performed for each salt/solvent species. AIMD training trajectories (NVT, "
            "Nosé–Hoover, POTIM=1 fs, 50,000 steps) at 300, 600, and 700 K generated the "
            "dataset for DeepMD model training (se_a descriptor, r_cut=6 Å, 500,000 steps, "
            "4×H100 GPUs, energy RMSE < 5 meV/atom). "
            "MLMD production runs used LAMMPS with the trained potential: NPT at 300 K "
            "for equilibration followed by NVT at each temperature for 1 ns. "
            "Self-diffusion coefficients were extracted from the MSD slope (40–80% lag window). "
            "Ionic conductivity was estimated via the Nernst-Einstein relation. "
            "Arrhenius activation energies were obtained by linear regression of ln(D) vs 1/T."
        )
        _insert_equation(doc, _eq_msd())
        _insert_equation(doc, _eq_diffusivity())
        _insert_equation(doc, _eq_nernst_einstein())
        _insert_equation(doc, _eq_arrhenius())

        # ── Results: comparative table ─────────────────────────────────────
        _h1("3. Results and Discussion")
        _h2("3.1 Comparative Transport Properties")
        _body(
            "Table 1 summarises the Arrhenius activation energy and electrochemical stability "
            "window for each composition."
        )

        has_echem = any(r.get("echem") for r in sub_results)
        headers = ["Composition", "Ea (eV)", "V_red (V)", "V_ox (V)"] if has_echem \
                  else ["Composition", "Ea (eV)"]
        rows = []
        for r in sub_results:
            row = [r["name"], r.get("Ea_str") or "—"]
            if has_echem:
                ec = r.get("echem") or {}
                row += [
                    f"{ec.get('V_red'):.2f}" if ec.get("V_red") is not None else "—",
                    f"{ec.get('V_ox'):.2f}"  if ec.get("V_ox")  is not None else "—",
                ]
            rows.append(row)
        _add_table(headers, rows)
        _body("Table 1. Comparative transport and electrochemical properties.", size=9.5)

        # ── Arrhenius figures ──────────────────────────────────────────────
        _h2("3.2 Arrhenius Analysis")
        _insert_equation(doc, _eq_activation_energy())
        for r in sub_results:
            _add_fig(r["fig_dir"], "arrhenius*.png",
                     f"Arrhenius plot for {r['name']}. Ea = {r.get('Ea_str', 'N/A')}.")

        # ── MSD figures ────────────────────────────────────────────────────
        _h2("3.3 Mean Square Displacement")
        for r in sub_results:
            _add_fig(r["fig_dir"], "msd*.png", f"MSD analysis for {r['name']}.")

        # ── RDF / Solvation structure ───────────────────────────────────────
        has_rdf = any(list((r["dir"] / "Analysis").glob("rdf_*.csv")) for r in sub_results)
        if has_rdf:
            _h2("3.4 Solvation Structure")
            _insert_equation(doc, _eq_rdf())
            _insert_equation(doc, _eq_cn())
            for r in sub_results:
                rdf_text = self._build_rdf_section(r["dir"])
                if rdf_text:
                    _body(f"{r['name']}: {rdf_text}")
                _add_fig(r["fig_dir"], "rdf*.png",
                         f"Radial distribution function for {r['name']}.")

        # ── Electronic structure ────────────────────────────────────────────
        has_elec = any((r["dir"] / "results" / "electronic" / "homo_lumo.json").exists()
                       for r in sub_results)
        if has_elec:
            _h2("3.5 Electronic Structure and Electrochemical Stability")
            _insert_equation(doc, _eq_koopmans_ip())
            _insert_equation(doc, _eq_koopmans_ea())
            _insert_equation(doc, _eq_vox())
            _insert_equation(doc, _eq_vred())
            for r in sub_results:
                sub_yaml = self.read_project_yaml(r["dir"])
                elec_text = self._build_electronic_section(r["dir"], sub_yaml)
                if elec_text:
                    _body(f"{r['name']}: {elec_text}")
                _add_fig(r["fig_dir"], "homo_lumo*.png",
                         f"HOMO/LUMO diagram for {r['name']}.")

        # ── Per-composition detailed sections ──────────────────────────────
        sec_n = 4
        _h1(f"{sec_n}. Per-Composition Detailed Results")
        for idx, r in enumerate(sub_results):
            _h2(f"{sec_n}.{idx + 1} {r['name']}")
            sub_yaml = self.read_project_yaml(r["dir"])
            _body(self._build_transport_section(r["dir"], sub_yaml))
            echem_json = r["dir"] / "results" / "echem" / "echem_summary.json"
            if echem_json.exists():
                _body(self._build_echem_section(r["dir"]))

        # ── Conclusions ────────────────────────────────────────────────────
        conc_sec = sec_n + 1
        _h1(f"{conc_sec}. Conclusions")
        Ea_list = [
            f"{r['name']}: {r['Ea_str']}" for r in sub_results if r.get("Ea_str")
        ]
        _body(
            f"We performed a systematic multi-scale computational study of {sol_str} "
            f"electrolytes at {len(sub_results)} concentrations. "
            "A unified DFT → AIMD → DeepMD → MLMD pipeline was applied consistently across "
            "all compositions, enabling direct comparison of transport properties. "
            + (f"Arrhenius activation energies: {'; '.join(Ea_list)}. " if Ea_list else "")
            + "Concentration-dependent solvation structure and electronic properties are "
            "characterised, providing molecular-level insight into the relationship between "
            "salt concentration and ionic transport in these electrolyte systems."
        )

        # ── References ─────────────────────────────────────────────────────
        _h1("References")
        for ref in yaml_cfg.get("references", self._default_references(yaml_cfg)):
            _body(ref, size=10)

        doc.save(str(out_path))
        log.info("[h11_manuscript] Combinatorial manuscript saved: %s", out_path)

    # ── Version management ───────────────────────────────────────────────────

    def _next_version(self, ms_dir: Path, name: str) -> int:
        """Return the next version number by scanning ms_dir and the optional global archive for existing vN files."""
        pattern = re.compile(
            rf"{re.escape(name)}_Manuscript_v(\d+)\.docx$", re.IGNORECASE
        )
        max_v = 0
        for search_dir in [ms_dir, _GLOBAL_MS_DIR]:
            if search_dir and search_dir.is_dir():
                for p in search_dir.glob("*.docx"):
                    m = pattern.match(p.name)
                    if m:
                        max_v = max(max_v, int(m.group(1)))
        return max_v + 1

    # ═══════════════════════════════════════════════════════════════════════
    # Generic DOCX generator — works for any project type
    # ═══════════════════════════════════════════════════════════════════════

    def _generate_generic_docx(self, project_dir: Path, yaml: dict,
                                out_path: Path) -> None:
        """Write a full single-project DOCX manuscript with adaptive sections based on available data."""
        if _CLADUE_SITE not in sys.path:
            sys.path.insert(0, _CLADUE_SITE)

        try:
            from docx import Document
            from docx.shared import Pt, Inches, Cm, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
        except ImportError:
            log.error("[h11_manuscript] python-docx not available")
            stub = self._collect_text_content(project_dir, yaml)
            out_path.with_suffix(".txt").write_text(stub)
            return

        project_name = yaml.get("name", project_dir.name)
        fig_dir = project_dir / "Analysis" / "figures"
        if not fig_dir.exists():
            fig_dir = project_dir / "analysis" / "figures"

        doc = Document()

        # Page margins
        section = doc.sections[0]
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

        # ── Style helpers ────────────────────────────────────────────────────
        def _run(para, text, bold=False, italic=False, size=11, font="Times New Roman"):
            """Add a styled run to para and return it."""
            r = para.add_run(text)
            r.bold = bold; r.italic = italic
            r.font.name = font; r.font.size = Pt(size)
            return r

        def _h1(text):
            """Append an underlined level-1 heading paragraph and return it."""
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after  = Pt(6)
            r = _run(p, text, bold=True, size=13)
            r.font.underline = True
            return p

        def _h2(text):
            """Append a bold level-2 heading paragraph and return it."""
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after  = Pt(4)
            _run(p, text, bold=True, size=11)
            return p

        def _body(text, size=11, indent_cm=0.0):
            """Append a body-text paragraph with optional left indent and return it."""
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(6)
            if indent_cm:
                p.paragraph_format.left_indent = Cm(indent_cm)
            _run(p, text, size=size)
            return p

        def _fig(name_stem, caption, width_in=5.5):
            """Insert the best available PNG for name_stem (with common suffixes), or a placeholder."""
            for suffix in ("", "_4panel", "_spline", "_orbital", "_enhanced"):
                cand = fig_dir / f"{name_stem}{suffix}.png"
                if cand.exists():
                    try:
                        doc.add_picture(str(cand), width=Inches(width_in))
                        cap = doc.add_paragraph()
                        _run(cap, caption, italic=True, size=9.5)
                        cap.paragraph_format.space_after = Pt(10)
                        return
                    except Exception:
                        pass
            _body(f"[Figure placeholder: {name_stem}.png]", size=9)

        def _table(headers, rows, col_widths=None):
            """Append a styled Word table with a blue header row and optional column widths."""
            tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
            tbl.style = "Table Grid"
            hrow = tbl.rows[0]
            for i, h in enumerate(headers):
                cell = hrow.cells[i]
                cell.text = ""
                r = cell.paragraphs[0].add_run(h)
                r.bold = True; r.font.size = Pt(9.5)
                r.font.color.rgb = RGBColor(255, 255, 255)
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "2E75B6")
                tcPr.append(shd)
            for ri, row in enumerate(rows):
                trow = tbl.rows[ri + 1]
                for ci, val in enumerate(row):
                    trow.cells[ci].text = ""
                    r = trow.cells[ci].paragraphs[0].add_run(str(val))
                    r.font.size = Pt(9.5)
            if col_widths:
                for ci, w in enumerate(col_widths):
                    for row in tbl.rows:
                        row.cells[ci].width = Inches(w)
            doc.add_paragraph()

        # ── Title, authors, affiliations ─────────────────────────────────────
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(title_p,
             yaml.get("manuscript_title",
                      f"Computational Study of {project_name}: "
                      "Ab Initio and Machine-Learning Molecular Dynamics Investigation"),
             bold=True, size=14)

        auth_p = doc.add_paragraph()
        auth_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(auth_p, yaml.get("authors", "Selva Chandrasekaran Selvaraj et al."), size=11)

        affil_p = doc.add_paragraph()
        affil_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(affil_p,
             yaml.get("affiliations",
                      "National Renewable Energy Laboratory (NREL), Golden, CO 80401"),
             italic=True, size=10)

        doc.add_paragraph()

        # ── Abstract ─────────────────────────────────────────────────────────
        _h1("Abstract")
        _body(yaml.get("abstract", self._build_abstract(project_dir, yaml)))

        # ── Introduction ──────────────────────────────────────────────────────
        _h1("1. Introduction")
        _body(yaml.get(
            "intro_text",
            f"This work presents a multiscale computational study of {project_name} "
            "combining ab initio molecular dynamics (AIMD), machine-learning interatomic "
            "potentials (MLIP), and machine-learning molecular dynamics (MLMD) to investigate "
            "structural, transport, and electronic properties relevant to next-generation "
            "electrochemical energy storage."
        ))

        # ── Computational Methods ─────────────────────────────────────────────
        _h1("2. Computational Methods")

        _h2("2.1 Density Functional Theory (DFT)")
        _body(self._build_dft_methods(yaml))

        # AIMD methods — if AIMD data present
        aimd_dir = project_dir / "aimd"
        if aimd_dir.is_dir() or any(project_dir.rglob("XDATCAR")):
            _h2("2.2 Ab Initio Molecular Dynamics (AIMD)")
            _body(self._build_aimd_methods(yaml))

        # MLIP methods — if pot_com.pb or MACE checkpoint present
        has_mlip = (any(project_dir.rglob("pot_com.pb")) or
                    any(project_dir.rglob("*.pth")))
        if has_mlip:
            _h2("2.3 Machine-Learning Interatomic Potential")
            _body(self._build_mlip_methods(yaml))

        # MLMD methods — if LAMMPS dump present
        has_mlmd = any(project_dir.rglob("dump_unwrapped.lmp"))
        if has_mlmd:
            _h2("2.4 Machine-Learning Molecular Dynamics (MLMD)")
            _body(
                "MLMD trajectories were generated with LAMMPS using the trained DeepMD "
                "potential (GPU backend, 4× H100). Simulations of 1 ns duration were "
                "performed at each temperature."
            )

        # ── Results: Transport ────────────────────────────────────────────────
        arrh_csv = self._find_arrhenius_csv(project_dir)
        has_msd  = any((project_dir / "Analysis").rglob("msd_*K.csv"))
        if has_msd or (arrh_csv is not None and arrh_csv.exists()):
            _h1("3. Results and Discussion")
            _h2("3.1 Li-Ion Diffusivity and Ionic Conductivity")
            _body(
                "The mean square displacement (MSD) was computed from MLMD trajectories "
                "at multiple temperatures:"
            )
            _insert_equation(doc, _eq_msd())
            _body(
                "The self-diffusion coefficient was extracted by linear regression of the "
                "MSD in the 40–80% lag-time window:"
            )
            _insert_equation(doc, _eq_diffusivity())
            _body(
                "Ionic conductivity was estimated from the Nernst-Einstein equation:"
            )
            _insert_equation(doc, _eq_nernst_einstein())
            _body(
                "An Arrhenius plot of ln(D) vs 1/T provides the activation energy:"
            )
            _insert_equation(doc, _eq_arrhenius())
            _insert_equation(doc, _eq_activation_energy())
            _body(self._build_transport_section(project_dir, yaml))

            # MSD figure
            if has_msd:
                _fig("msd_4panel",
                     "Figure. Mean square displacement (4-panel). "
                     "(a) MSD vs time (raw + Savitzky-Golay smoothed). "
                     "(b) log-log MSD. (c) β exponent. (d) D(t) = d(MSD)/dt / 6.")

            # Arrhenius figure
            if arrh_csv is not None and arrh_csv.exists():
                _fig("arrhenius",
                     "Figure. Arrhenius plot: ln(D) vs 1000/T. "
                     "Open symbols — MLMD data; dashed line — linear fit. "
                     f"Ea = {self._read_ea(arrh_csv)} eV.")

        # ── Results: Electronic structure ─────────────────────────────────────
        homo_lumo = project_dir / "results" / "electronic" / "homo_lumo.json"
        bader_csv = project_dir / "results" / "electronic" / "bader_charges.csv"
        dos_csv   = project_dir / "results" / "electronic" / "dos_total.csv"

        # Initialize section counter before the block so echem/RDF can reference
        # sec_n and _sub even when no electronic structure data is present.
        sec_n = "4" if (has_msd or arrh_csv.exists()) else "3"
        _sub = [0]
        def _next_sub():
            """Return the next sub-section label (e.g. '4.2') and increment the counter."""
            _sub[0] += 1
            return f"{sec_n}.{_sub[0]}"

        if homo_lumo.exists() or bader_csv.exists() or dos_csv.exists():
            _h1(f"{sec_n}. Electronic Structure")

            if bader_csv.exists():
                _h2(f"{_next_sub()} Bader Charge Analysis")
                _body(
                    "Bader charge analysis (Henkelman method) reveals the charge transfer "
                    "between the mobile ion and the host framework."
                )
                _fig("bader_charges", "Figure. Bader charge transfer per atom.")

            if dos_csv.exists():
                _h2(f"{_next_sub()} Density of States")
                _body(
                    "The total and projected density of states (DOS/PDOS) characterise the "
                    "electronic structure near the Fermi level."
                )
                _fig("dos_total", "Figure. Total density of states (DOS). "
                     "Vertical dashed line at E_F = 0.")

            if homo_lumo.exists():
                _h2(f"{_next_sub()} HOMO/LUMO and Electrochemical Stability")
                _body(
                    "Frontier orbital energies are connected to electrochemical stability "
                    "via Koopmans' theorem (Koopmans, 1934):"
                )
                _insert_equation(doc, _eq_koopmans_ip())
                _insert_equation(doc, _eq_koopmans_ea())
                _body(
                    "Conversion to the Li/Li⁺ electrochemical scale uses the Trasatti (1986) "
                    "absolute electrode potential (Li/Li⁺ vs vacuum = 1.37 eV):"
                )
                _insert_equation(doc, _eq_vox())
                _insert_equation(doc, _eq_vred())
                _body(self._build_electronic_section(project_dir, yaml))
                _fig("homo_lumo_orbital",
                     "Figure. HOMO/LUMO energy level diagram (DFT-PBE, vacuum reference). "
                     "Filled bars: HOMO. Dashed bars: LUMO. Gap values labelled. "
                     "Dotted line: Li/Li⁺ reference (−1.37 eV vs vacuum).")

        # ── Results: Electrochemical stability window ─────────────────────────
        echem_json = project_dir / "results" / "echem" / "echem_summary.json"
        if echem_json.exists():
            _h2(f"{sec_n}.{_sub[0] + 1} Electrochemical Stability Windows")
            _sub[0] += 1
            _body(self._build_echem_section(project_dir))
            _fig("echem_window",
                 "Figure. Electrochemical stability windows (V_red to V_ox). "
                 "Shaded band: typical Li-ion operating window (0–4.2 V).")

        # ── Results: Solvation / RDF ──────────────────────────────────────────
        rdf_files = list((project_dir / "Analysis").glob("rdf_*.csv"))
        if rdf_files:
            _h2(f"{sec_n}.{_sub[0] + 1} Solvation Structure — Radial Distribution Functions")
            _sub[0] += 1
            _body(
                "The radial distribution function g(r) describes the local structure "
                "around the mobile ion:"
            )
            _insert_equation(doc, _eq_rdf())
            _body("Coordination numbers CN are obtained by integrating to the first minimum r_min:")
            _insert_equation(doc, _eq_cn())
            _body(self._build_rdf_section(project_dir))
            for rdf_f in sorted(rdf_files)[:3]:
                pair = rdf_f.stem.replace("rdf_", "")
                # Strip temperature suffix (e.g. "_300K") to get the figure stem
                pair_stem = re.sub(r"_\d+K$", "", pair)
                _fig(f"rdf_{pair_stem}",
                     f"Figure. RDF g(r) for {pair_stem} pair at 300 K. "
                     "Right axis: running coordination number N(r).")

        # ── Results: NEB migration barriers ──────────────────────────────────
        # Compute NEB section number dynamically based on what precedes it
        neb_json = project_dir / "results" / "neb_barriers.json"
        if neb_json.exists():
            _neb_sec = int(sec_n) + 1  # one major section after electronic/results
            _h1(f"{_neb_sec}. NEB Migration Barriers")
            _body(
                "Nudged elastic band (CI-NEB) calculations provide kinetic barriers "
                "for ion migration or bond-breaking events:"
            )
            _insert_equation(doc, _eq_neb_barrier())
            _body(self._build_neb_section(project_dir))
            _fig("neb_profile_spline",
                 "Figure. NEB energy profiles (cubic spline interpolation). "
                 "IS = initial state, TS = transition state (CI image), FS = final state. "
                 "Ea = activation barrier, ΔE = reaction energy.")
            _fig("neb_barriers",
                 "Figure. NEB migration barrier summary (meV).")
        else:
            _neb_sec = int(sec_n)

        # ── Atomic structures ─────────────────────────────────────────────────
        if dft_opt(project_dir).is_dir():
            _h1(f"{_neb_sec + 1}. Atomic Structures")
            _body(
                "DFT-PBE optimised structures are shown below "
                "(CPK colour scheme; H=light grey, C=charcoal, O=red, N=blue, Li=violet, "
                "Na=dark blue, S=gold, F=cyan, Cl=green, Mg=pink, Zn=brown)."
            )
            _fig("atomic_structures_panel",
                 "Figure. Ball-and-stick representations of DFT-optimised structures "
                 f"for {project_dir.name}.")

        # ── Conclusions ─ section number adapts to how many sections precede it ──
        _has_opt = dft_opt(project_dir).is_dir()
        _conc_sec = _neb_sec + 1 + (1 if _has_opt else 0)
        _h1(f"{_conc_sec}. Conclusions")
        _body(yaml.get(
            "conclusion_text",
            f"We have performed a multiscale computational study of {project_name}. "
            "DFT-AIMD provides training data for a DeepMD machine-learning potential, "
            "enabling long-timescale MLMD simulations of ionic transport. "
            "The computed activation energy and room-temperature conductivity are "
            "reported. Electronic structure analysis and NEB migration barriers "
            "characterise the material's stability and ion-transport mechanism."
        ))

        # ── Data Availability ─────────────────────────────────────────────────
        _h1("Data Availability")
        _body(
            "All VASP input/output files, LAMMPS trajectories, and analysis CSV files "
            "are available in the project repository at "
            f"{project_dir}."
        )

        # ── References ────────────────────────────────────────────────────────
        _h1("References")
        refs = yaml.get("references", self._default_references(yaml))
        for ref in refs:
            _body(ref, size=10)

        doc.save(str(out_path))
        log.info("[h11_manuscript] Saved %s", out_path)

    # ── Content builders ─────────────────────────────────────────────────────

    def _build_abstract(self, project_dir: Path, yaml: dict) -> str:
        """Compose an abstract paragraph from available MSD temperatures and Arrhenius Ea."""
        project_name = yaml.get("name", project_dir.name)
        _arrh_csv = self._find_arrhenius_csv(project_dir)
        Ea_str    = (self._read_ea(_arrh_csv) if _arrh_csv else None) or "N/A"
        temps_str = ", ".join(
            str(_extract_T(p.name)) + " K"
            for p in sorted((project_dir / "Analysis").glob("msd_*K.csv"),
                            key=lambda x: _extract_T(x.name))
        ) or "multiple temperatures"

        return (
            f"We present a first-principles and machine-learning computational study of "
            f"{project_name} for next-generation electrochemical energy storage applications. "
            "Ab initio molecular dynamics (AIMD, VASP 6.4.2, PBE-GGA) at multiple temperatures "
            "provides training data for a DeepMD machine-learning interatomic potential. "
            f"Machine-learning molecular dynamics (MLMD, LAMMPS) at {temps_str} yields "
            "mean-square-displacement analysis and Nernst-Einstein ionic conductivity estimates. "
            f"The Arrhenius activation energy for ion diffusion is Ea = {Ea_str}. "
            "Electronic structure calculations (Bader charges, DOS, HOMO-LUMO) characterise "
            "the material's stability. NEB migration barriers quantify the kinetic barriers "
            "for ion transport and decomposition pathways."
        )

    def _build_dft_methods(self, yaml: dict) -> str:
        """Return a DFT methods paragraph filled from project.yaml VASP settings."""
        encut  = yaml.get("encut",        520.0)
        xc     = yaml.get("xc_functional", "PBE")
        kpts   = yaml.get("kpoints",       [2, 2, 2])
        ediff  = yaml.get("ediff",         "1E-5")
        ediffg = yaml.get("ediffg",        "-0.02")
        return (
            f"All DFT calculations were performed with VASP 6.4.2 using PAW pseudopotentials "
            f"and the {xc} exchange-correlation functional. "
            f"Plane-wave kinetic energy cutoff: {encut} eV. "
            f"Brillouin-zone sampling: {kpts[0]}×{kpts[1]}×{kpts[2]} Γ-centred k-mesh. "
            f"Electronic convergence: EDIFF = {ediff} eV; "
            f"Ionic convergence: EDIFFG = {ediffg} eV/Å. "
            "Dispersion correction: DFT-D3 (Grimme, Becke-Johnson damping, IVDW=12). "
            "ISYM=0 (no symmetry enforcement in MD/NEB). LREAL=Auto."
        )

    def _build_aimd_methods(self, yaml: dict) -> str:
        """Return an AIMD methods paragraph with ensemble, temperatures, and step count from project.yaml."""
        temps  = yaml.get("aimd_temps", [300, 600, 700])
        nsw    = yaml.get("nsw_aimd",   50000)
        dt_fs  = yaml.get("potim_fs",   1.0)
        return (
            f"AIMD simulations used the NVT ensemble (Nosé–Hoover thermostat, SMASS=0) "
            f"at {', '.join(str(T) + ' K' for T in temps)}. "
            f"Time step: {dt_fs} fs; {nsw} steps per temperature "
            f"(total: {nsw * dt_fs / 1000:.1f} ps per run). "
            "AIMD INCAR: ENCUT=400.8 eV, EDIFF=1E-4, NELM=60, IBRION=0, LREAL=Auto."
        )

    def _build_mlip_methods(self, yaml: dict) -> str:
        """Return a DeepMD training methods paragraph with descriptor and training settings from project.yaml."""
        rcut   = yaml.get("deepmd_rcut",   6.0)
        neuron = yaml.get("deepmd_neuron", "[25,50,100]")
        nsteps = yaml.get("deepmd_nsteps", 500000)
        return (
            "A Deep Potential Molecular Dynamics (DeepMD-kit) model was trained on "
            f"AIMD frames using the se_a descriptor (r_cut = {rcut} Å, "
            f"neuron = {neuron}, embedding/fitting nets, {nsteps} training steps). "
            "Training was performed on NREL Kestrel (4× NVIDIA H100 GPUs). "
            "Acceptance criteria: energy RMSE < 5 meV/atom, force RMSE < 100 meV/Å. "
            "Model frozen (pot.pb) and compressed (pot_com.pb) before production MD."
        )

    def _build_transport_section(self, project_dir: Path, yaml: dict) -> str:
        """Return a 2-sentence transport results paragraph with Ea and D extraction method."""
        _arrh_csv = self._find_arrhenius_csv(project_dir)
        Ea_str = (self._read_ea(_arrh_csv) if _arrh_csv else None) or "not computed"
        mobile_ion = yaml.get("mobile_ion", "Li")
        return (
            f"The Arrhenius activation energy for {mobile_ion}-ion diffusion is "
            f"Ea = {Ea_str}. "
            "Self-diffusivity D was extracted from the MSD slope in the 40–80% lag-time "
            "window (linear regime where β ≈ 1). Conversion: D [m²/s] = slope [Å²/ps] × 10⁻⁸."
        )

    def _build_electronic_section(self, project_dir: Path, yaml: dict) -> str:
        """Return a text summary of HOMO/LUMO energies and electrochemical windows from homo_lumo.json."""
        homo_lumo_json = project_dir / "results" / "electronic" / "homo_lumo.json"
        if not homo_lumo_json.exists():
            return "Electronic structure analysis results are pending."
        try:
            data = json.loads(homo_lumo_json.read_text())
            mols = data if isinstance(data, list) else list(data.values())
            names = list(data.keys()) if isinstance(data, dict) else []
            lines = []
            for name, mol in zip(names, mols):
                h = mol.get("HOMO", -5.5)
                l = mol.get("LUMO", -0.5)
                vox  = -h - 1.37
                vred = -l - 1.37
                lines.append(f"{name}: HOMO={h:.3f} eV, LUMO={l:.3f} eV, "
                             f"gap={l-h:.3f} eV, V_ox={vox:.2f} V, V_red={vred:.2f} V.")
            return " ".join(lines)
        except Exception:
            return "Frontier orbital energies from DFT-PBE (HOMO/LUMO diagram above)."

    def _build_echem_section(self, project_dir: Path) -> str:
        """Return a text summary of electrochemical stability windows from echem_summary.json."""
        echem_json = project_dir / "results" / "echem" / "echem_summary.json"
        try:
            d = json.loads(echem_json.read_text())
            lines = []
            # Handle species-list format (from Koopmans / HOMO-LUMO data)
            if "species" in d and isinstance(d["species"], list):
                for sp in d["species"]:
                    name = sp.get("name", "?")
                    vred = sp.get("V_red", 0.0)
                    vox  = sp.get("V_ox",  0.0)
                    win  = sp.get("window_V", vox - vred)
                    lines.append(
                        f"{name}: V_red = {vred:.2f} V, V_ox = {vox:.2f} V, "
                        f"window = {win:.2f} V."
                    )
                return " ".join(lines)
            # Handle flat format
            if "window_V" in d:
                lines.append(f"Electrochemical stability window: "
                             f"{d.get('V_red', '?'):.3f}–{d.get('V_ox', '?'):.3f} V "
                             f"(width = {d['window_V']:.3f} V).")
            if "OCV_V" in d:
                lines.append(f"Open-circuit voltage: {d['OCV_V']:.3f} V.")
            if "formation_energy_eV_per_atom" in d:
                lines.append(f"Formation energy: {d['formation_energy_eV_per_atom']:.4f} eV/atom.")
            return " ".join(lines) if lines else "Electrochemical stability computed by Koopmans' theorem."
        except Exception:
            return "Electrochemical properties computed from DFT."

    def _build_rdf_section(self, project_dir: Path) -> str:
        """Return a text description of first-peak positions from rdf_*.csv files."""
        rdf_files = sorted((project_dir / "Analysis").glob("rdf_*.csv"))
        if not rdf_files:
            return ""
        try:
            import numpy as np
            lines = []
            for f in rdf_files[:4]:
                pair = re.sub(r"_\d+K$", "", f.stem.replace("rdf_", ""))
                data = np.loadtxt(str(f), delimiter=",", skiprows=1)
                if data.ndim > 1 and data.shape[0] > 5:
                    r1 = data[data[:, 1].argmax(), 0]
                    lines.append(f"{pair}: first peak at r₁ ≈ {r1:.2f} Å.")
            return " ".join(lines)
        except Exception:
            return "Radial distribution functions characterise the solvation/coordination shell."

    def _build_neb_section(self, project_dir: Path) -> str:
        """Return a text summary of NEB barriers and mechanisms from neb_barriers.json."""
        neb_json = project_dir / "results" / "neb_barriers.json"
        try:
            neb_data = json.loads(neb_json.read_text())
            lines = []
            for path_name, pdata in neb_data.items():
                Ea = pdata.get("Ea_meV", pdata.get("Ea_eV", "?"))
                unit = "meV" if "Ea_meV" in pdata else "eV"
                mech = pdata.get("mechanism", "")
                lines.append(f"{path_name}: Ea = {Ea} {unit} ({mech}).")
            return " ".join(lines)
        except Exception:
            return "NEB barriers extracted from CI-NEB calculations (VASP 6.4.2, IMAGES=4–11, LCLIMB=.TRUE.)."

    def _collect_text_content(self, project_dir: Path, yaml: dict) -> str:
        """Return a plain-text manuscript stub used as fallback when python-docx is unavailable."""
        project_name = yaml.get("name", project_dir.name)
        return (
            f"Manuscript stub for {project_name}\n"
            "Generated by h11_manuscript.py (python-docx not available)\n\n"
            f"Abstract:\n{self._build_abstract(project_dir, yaml)}\n\n"
            f"DFT Methods:\n{self._build_dft_methods(yaml)}\n"
        )

    def _default_references(self, yaml: dict) -> list[str]:
        """Return the standard reference list, appending Koopmans/Trasatti refs for electrolyte projects."""
        refs = [
            "[1] Kresse, G.; Furthmüller, J. Phys. Rev. B 1996, 54, 11169. (VASP)",
            "[2] Kresse, G.; Joubert, D. Phys. Rev. B 1999, 59, 1758. (PAW)",
            "[3] Perdew, J. P.; Burke, K.; Ernzerhof, M. Phys. Rev. Lett. 1996, 77, 3865. (PBE)",
            "[4] Grimme, S. et al. J. Comput. Chem. 2011, 32, 1456. (DFT-D3)",
            "[5] Wang, H. et al. Comput. Phys. Commun. 2018, 228, 178. (DeepMD-kit)",
            "[6] Thompson, A. P. et al. Comput. Phys. Commun. 2022, 271, 108171. (LAMMPS)",
            "[7] Ong, S. P. et al. Comput. Mater. Sci. 2013, 68, 314. (pymatgen)",
            "[8] Larsen, A. H. et al. J. Phys.: Condens. Matter 2017, 29, 273002. (ASE)",
            "[9] Henkelman, G. et al. J. Chem. Phys. 2000, 113, 9901. (CI-NEB)",
            "[10] Henkelman, G.; Arnaldsson, A.; Jonsson, H. Comput. Mater. Sci. 2006, 36, 354. (Bader)",
        ]
        project_type = yaml.get("type", "").lower()
        if "electrolyte" in project_type or "solvent" in project_type:
            refs += [
                "[11] Koopmans, T. Physica 1934, 1, 104. (Koopmans' theorem)",
                "[12] Trasatti, S. Pure Appl. Chem. 1986, 58, 955. (Absolute electrode potential)",
            ]
        return refs

    @staticmethod
    def _find_arrhenius_csv(project_dir: Path) -> Path | None:
        """Return the best arrhenius.csv path under Analysis/{variant}/, preferring aimd."""
        aimd_csv = project_dir / "Analysis" / "aimd" / "arrhenius.csv"
        if aimd_csv.exists():
            return aimd_csv
        candidates = sorted(project_dir.glob("Analysis/*/arrhenius.csv"))
        return candidates[0] if candidates else None

    @staticmethod
    def _read_ea(arrh_csv: Path) -> str | None:
        """Read Ea (eV) from column 5 of arrhenius.csv; returns formatted string or None."""
        if not arrh_csv.exists():
            return None
        try:
            import numpy as np
            data = np.loadtxt(str(arrh_csv), delimiter=",", skiprows=1)
            if data.ndim == 1:
                data = data.reshape(1, -1)
            if data.shape[1] > 4:
                return f"{data[0, 4]:.3f} eV"
        except Exception:
            pass
        return None


def _extract_T(filename: str) -> int:
    """Extract temperature integer from a filename like 'msd_600K.csv'; returns 0 if not found."""
    m = re.search(r"(\d+)K", filename)
    return int(m.group(1)) if m else 0
